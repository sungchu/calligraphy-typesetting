import re
import time
import random
import streamlit as st
from streamlit_image_select import image_select
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.core.os_manager import ChromeType
from collections import defaultdict
import os
from zhconv import convert
from PIL import Image
import requests
from io import BytesIO

# 設定頁面寬度
st.set_page_config(layout="wide")

# ================= 初始化 session_state =================
if "results" not in st.session_state:
    st.session_state.results = []
if "selected_images" not in st.session_state:
    st.session_state.selected_images = []
if "display_index" not in st.session_state:
    st.session_state.display_index = {}

st.title("書法字典圖片瀏覽器")
st.markdown(
    """
    <style>
    .stImage img {
        width: 120px !important;   /* 強制寬度 */
        height: auto !important;   /* 保持比例 */
        border-radius: 0 !important; /* 順便把圓角也拿掉 */
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ================= 輸入欄位 =================
col_input, col_style, col_dl, col_calligrapher = st.columns([2,0.8,0.8,1.8])
with col_input:
    search_input = st.text_input("請輸入要搜尋的文字（標點符號將自動忽略，可同時輸入多個字，建議長度不超過30字）")
    search_input_chinese = "".join(re.findall(r"[\u4e00-\u9fff]+", search_input))
with col_style:
    style_dict = {"1": "章草", "3": "篆書", "4": "簡牘", "5": "魏碑",
                  "6": "隸書", "7": "草書", "8": "行書", "9": "楷書"}
    style_value = st.selectbox("選擇書法字體",
                               options=list(style_dict.keys()),
                               format_func=lambda x: style_dict[x],
                               index=7)
with col_dl:
    # download_limit = 6
    download_limit = st.number_input("每個字最多出現幾個選項",min_value=1, max_value=10, value=5, step=1)
with col_calligrapher:
    filter_calligrapher_input = st.text_input(
        "指定特定書法家（若想指定多位，請用、分隔，留空則代表不指定 e.g. 王羲之、顏真卿）", ""
    )
    if filter_calligrapher_input.strip():
        filter_calligrapher_list = [c.strip() for c in filter_calligrapher_input.split("、") if c.strip()]
    else:
        filter_calligrapher_list = None


placeholder_img_path = os.path.join(os.getcwd(), "查無此字.png")  # 同資料夾下

# ================= 安全顯示圖片 =================
def safe_show_image(img_url, width=5):
    try:
        if not img_url:
            img = Image.open(placeholder_img_path)
            st.image(img, width=width)
            return

        if isinstance(img_url, str) and img_url.startswith("http"):
            resp = requests.get(img_url, timeout=5)
            if resp.status_code != 200:
                st.write("⚠️ 圖片下載失敗")
                return
            img = Image.open(BytesIO(resp.content))
            st.image(img, width=width)
        else:
            # 假設是本地檔案
            if os.path.exists(img_url):
                img = Image.open(img_url)
                st.image(img, width=width)
            else:
                st.write("⚠️ 找不到圖片檔案")
    except Exception as e:
        st.write(f"⚠️ 無法顯示圖片：{e}")

# ================= 預覽word內容 =================
import streamlit as st
import requests
from PIL import Image
from io import BytesIO
import base64

def image_to_base64(img_url, width=120):
    """下載圖片並轉成 base64 方便在 HTML table 顯示"""
    try:
        if img_url and img_url.startswith("http"):
            resp = requests.get(img_url, timeout=5)
            resp.raise_for_status()
            img = Image.open(BytesIO(resp.content))
        else:
            return None
        # 縮小圖片
        img.thumbnail((width, width))
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        b64 = base64.b64encode(buffer.getvalue()).decode()
        return f"data:image/png;base64,{b64}"
    except Exception:
        return None


def preview_layout(selected_data):
    """在 Streamlit 畫出和 Word 相同的 4直行×12橫排 排版"""
    sorted_selected = sorted(selected_data, key=lambda x: x[0])

    # 建立交替資料：word → image → word → image
    layout_items = []
    for _, word, author, img_url in sorted_selected:
        layout_items.append(("word", word[0]))
        layout_items.append(("image", img_url))

    total_cells = 12 * 4
    cells = [["" for _ in range(4)] for _ in range(12)]

    for idx, item in enumerate(layout_items[:total_cells]):
        row = idx % 12
        col = 3 - (idx // 12)  # 右到左
        if item[0] == "word":
            cells[row][col] = f"<div style='font-size:20px;text-align:center'>{item[1]}</div>"
        else:
            img_b64 = image_to_base64(item[1], width=60)
            if img_b64:
                cells[row][col] = f"<img src='{img_b64}' style='display:block;margin:auto;width:60px;height:auto;'/>"
            else:
                cells[row][col] = "<div style='color:red;text-align:center'>[圖片失敗]</div>"

    # 轉成 HTML table
    table_html = "<table style='border-collapse:collapse;margin:auto;'>"
    for r in range(12):
        height = "30px" if (r+1) % 2 == 1 else "90px"  # 模擬 Word 行高
        table_html += f"<tr style='height:{height};'>"
        for c in range(4):
            table_html += f"<td style='border:1px solid #ccc;width:120px;text-align:center;vertical-align:middle'>{cells[r][c]}</td>"
        table_html += "</tr>"
    table_html += "</table>"

    st.markdown(table_html, unsafe_allow_html=True)

# ================= 下載word =================
from docx import Document
from docx.shared import Cm, Pt
from io import BytesIO
import requests
from PIL import Image

def download_word(selected_data):
    """
    selected_data: list of tuples (idx, word, author, img_url)
    """

    # 建立 Word 文件
    doc = Document()
    
    # 設定頁面邊界：上下 1 公分，左右可依需求設定
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    # 如果也要左右設定，可以加上：
    # section.left_margin = Cm(1)
    # section.right_margin = Cm(1)
    
    # 建立表格：12 橫排、4 直行
    table = doc.add_table(rows=12, cols=5)
    table.autofit = False

    # 設定行高：單數行高 1、偶數行高 3
    for i, row in enumerate(table.rows):
        if (i+1) % 2 == 1:  # 單數行
            row.height = Cm(0.8)
        else:  # 偶數行
            row.height = Cm(3)

    # 資料排序（從右上往下）
    sorted_selected = sorted(selected_data, key=lambda x: x[0])

    # 建立交替資料：word → image → word → image
    layout_items = []
    for _, word, author, img_url in sorted_selected:
        layout_items.append(("word", word[0]))
        layout_items.append(("image", img_url))

    # 從右上角開始填（先列後行）
    total_cells = 12 * 5
    for idx, item in enumerate(layout_items[:total_cells]):
        row = idx % 12  # 行
        col = 4 - (idx // 12)  # 右到左的列（0→左,4→右）
        
        if col < 0 or col > 4 or row < 0 or row > 11:
            continue  # 超界就跳過
        
        cell = table.cell(row, col)

        if item[0] == "word":
            # 插入文字
            p = cell.paragraphs[0]
            run = p.add_run(item[1])
            run.font.size = Pt(12)
        else:
            # 插入圖片
            img_url = item[1]
            try:
                response = requests.get(img_url)
                image = Image.open(BytesIO(response.content))
                image_stream = BytesIO()
                image.save(image_stream, format="PNG")
                image_stream.seek(0)
                cell.paragraphs[0].add_run().add_picture(image_stream, width=Cm(3))
            except Exception as e:
                cell.text = "[圖片載入失敗]"

    # 存到記憶體
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ================= 搜尋按鈕 =================
if st.button("開始搜尋"):
    st.session_state.results = []
    st.session_state.selected_images = []
    st.session_state.display_index = {}

    search_words = list(search_input_chinese.strip())
    results = []
    total_words = len(search_words)
    progress_bar = st.progress(0, text='搜尋中，請稍後')
    status_text = st.empty()
    
    # Spinner 轉圈圈
    with st.spinner("🔄 搜尋中，請稍後..."):
        # ================= Selenium 設定 =================
        options = Options()
        options.add_argument("--headless")
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        options.add_argument("--incognito")
        options.add_argument("--disable-dev-shm-usage")
        service = Service(ChromeDriverManager(chrome_type=ChromeType.CHROMIUM).install())

        driver = None
        try:
            driver = webdriver.Chrome(service=service, options=options)
            base_url = "https://www.shufazidian.com/s.php"
            start_time = time.time()

            for idx, word in enumerate(search_words):
                word_found = False
                try:
                    driver.get(base_url)
                    time.sleep(random.uniform(1, 2))

                    search_input_elem = driver.find_element(By.ID, "wd")
                    search_input_elem.clear()
                    search_input_elem.send_keys(word)

                    select = Select(driver.find_element(By.ID, "sort"))
                    select.select_by_value(style_value)

                    submit_button = driver.find_element(By.XPATH, "//form[@name='form1']//button[@type='submit']")
                    submit_button.click()
                    time.sleep(random.uniform(2, 3))

                    j_elements = driver.find_elements(By.CSS_SELECTOR, "div.j")
                    if not filter_calligrapher_list:
                        j_elements = j_elements[1:]

                    for j in j_elements[:-1]:
                        try:
                            a_img = j.find_element(By.CSS_SELECTOR, "div.mbpho a")
                            img_url = a_img.get_attribute("href")
                        except:
                            continue

                        try:
                            g_div = j.find_element(By.CSS_SELECTOR, "div.g")
                            a_btnsfj = g_div.find_element(By.CSS_SELECTOR, "a.btnSFJ")
                            author_name = a_btnsfj.get_attribute("sfj") or a_btnsfj.text.strip()
                        except:
                            author_name = g_div.text.split('\n')[0].strip()

                        if filter_calligrapher_list and (convert(author_name,'zh-tw') not in filter_calligrapher_list):
                            continue

                        results.append((word, author_name, img_url))
                        word_found = True

                    if not word_found:
                        if not any(r[0] == word and r[1] == "查無此字" for r in results):
                            results.append((word, "查無此字", None))

                except Exception as e:
                    results.append((word, "查無此字", None))
                    st.warning(f"{word} 搜尋失敗: {e}")

                completed = idx + 1
                elapsed = time.time() - start_time
                avg_time = elapsed / completed
                remaining = (total_words - completed) * avg_time
                remain_min = int(remaining // 60)
                remain_sec = int(remaining % 60)

                progress_bar.progress(completed / total_words)
                status_text.text(
                    f"正在搜尋第 {completed}/{total_words} 個字：{word} ⏳ 預估剩餘 {remain_min}分{remain_sec}秒"
                )

        finally:
            if driver:
                driver.quit()

        # 如果整體沒有任何圖片，給 placeholder
        has_any_image = any(img_url for _, _, img_url in results)
        if not has_any_image:
            results = [(word, "查無此字", placeholder_img_path) for word in search_words]

        st.session_state.results = results

        # 初始化 display_index，字+instance_id
        word_count = defaultdict(int)
        for word in search_words:
            word_count[word] += 1
            instance_id = word_count[word]
            st.session_state.display_index[f"{word}_{instance_id}"] = 0

col_1, col_select, col_empty, col_show, col4= st.columns([0.2, 0.9, 0.05, 1, 0.2])
# ================= 顯示搜尋結果 & 下一批圖片功能 =================
with col_select:
    results = st.session_state.get("results", [])
    if results:
        search_words = list(search_input_chinese.strip())
        groups_dict = defaultdict(list)
        for word, author, img_url in results:
            groups_dict[word].append((word, author, img_url))

        word_count = defaultdict(int)
        for w_idx, w in enumerate(search_words):
            word_count[w] += 1
            instance_id = word_count[w]
            group_items = groups_dict.get(w, [])
            if not group_items:
                continue

            st.subheader(f"🔍 {w} ({style_dict[style_value]})")
            start = st.session_state.display_index.get(f"{w}_{instance_id}", 0)
            end = min(start + download_limit, len(group_items))
            batch_items = group_items[start:end]

            # 確保每個 batch 至少有一張圖片
            img_urls = [img_url if img_url else placeholder_img_path for _, _, img_url in batch_items]
            labels = [convert(author,'zh-tw') if img_url else "查無此字" for _, author, img_url in batch_items]

            # 假設自動選第一張圖片，避免下一批按鈕要點圖片
            auto_select_first = st.session_state.get(f"auto_select_first_{w}_{instance_id}", True)
            default_index = 0 if auto_select_first and batch_items else None

            selected_idx = image_select(
                label=f"選擇 {w} 的圖片",
                images=img_urls,
                captions=labels,
                return_value="index",
                use_container_width=False,
                key=f"img_select_{w}_{instance_id}_{start}"
            )

            # 限制每組只能選一張圖片
            st.session_state.selected_images = [
                x for x in st.session_state.selected_images if x[1] != f"{w}_{instance_id}"
            ]
            if selected_idx is not None:
                word_sel, author_name, img_url = batch_items[selected_idx]
                st.session_state.selected_images.append((w_idx, f"{w}_{instance_id}", author_name, img_url))

            # 下一批按鈕
            if end < len(group_items):
                if st.button(f"下一批 {w}", key=f"next_batch_{w}_{instance_id}"):
                    # 更新顯示索引
                    st.session_state.display_index[f"{w}_{instance_id}"] = start + download_limit
                    # 自動選第一張
                    st.session_state[f"auto_select_first_{w}_{instance_id}"] = True
                    # 重新渲染頁面
                    st.rerun()
                else:
                    st.session_state[f"auto_select_first_{w}_{instance_id}"] = False
    # ================= 顯示挑選圖片 =================                
with col_show:
    # 建立一個 placeholder 用來顯示更新狀態
    status_placeholder = st.empty()

    if st.session_state.selected_images:
        # 顯示更新中
        status_placeholder.text("⏳ 更新中…")
        
        # 顯示預覽表格
        preview_layout(st.session_state.selected_images)

        # 下載 Word
        buffer = download_word(st.session_state.selected_images)
        st.download_button(
            label="📥 下載 Word",
            data=buffer,
            file_name="selected_images.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # 清掉更新文字
        status_placeholder.empty()
